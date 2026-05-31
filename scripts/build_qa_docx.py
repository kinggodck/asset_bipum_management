from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "ASSET_QR_MANAGER_QA.md"
OUTPUT = ROOT / "ASSET_QR_MANAGER_QA.docx"

FONT = "Malgun Gothic"
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(91, 103, 117)
LIGHT_FILL = "E8EEF5"
CODE_FILL = "F4F6F9"
BORDER = "D8DEE8"


def set_east_asia_font(run, font_name=FONT):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def set_paragraph_border(paragraph, color=BORDER, size="6", space="3"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def shade_paragraph(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_twips):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER)
        borders.append(tag)


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("비품 QR 관리 프로그램 Q&A")
    set_east_asia_font(run)
    run.font.size = Pt(8)
    run.font.color.rgb = MUTED


def add_title(doc):
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    run = title.add_run("비품 QR 관리 프로그램 Q&A")
    set_east_asia_font(run)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(23, 32, 51)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("직원 안내 및 운영 전환 참고 문서")
    set_east_asia_font(run)
    run.font.size = Pt(11)
    run.font.color.rgb = MUTED
    set_paragraph_border(subtitle)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(4)
    note.paragraph_format.space_after = Pt(12)
    shade_paragraph(note, LIGHT_FILL)
    run = note.add_run(
        "이 문서는 현재 비품 QR 관리 앱의 동작 방식, 저장 위치, Google Sheet 연동 한계, "
        "사진/QR/사용자 추적 관련 운영 질문을 직원용 Q&A 형태로 정리한 자료입니다."
    )
    set_east_asia_font(run)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 58, 95)


def parse_inline_code(paragraph, text):
    parts = re.split(r"(`[^`]+`)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            set_east_asia_font(run, "Consolas")
            run.font.size = Pt(9.5)
            run.font.color.rgb = DARK_BLUE
        else:
            run = paragraph.add_run(part)
            set_east_asia_font(run)


def add_code_block(doc, lines):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(3)
    paragraph.paragraph_format.space_after = Pt(8)
    shade_paragraph(paragraph, CODE_FILL)
    run = paragraph.add_run("\n".join(lines))
    set_east_asia_font(run, "Consolas")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(17, 24, 39)


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    parse_inline_code(paragraph, text)


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    parse_inline_code(paragraph, text)


def add_question_index(doc, questions):
    doc.add_heading("질문 목록", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.autofit = False
    set_table_borders(table)
    widths = [900, 8460]
    hdr = table.rows[0].cells
    hdr[0].text = "번호"
    hdr[1].text = "질문"
    for idx, cell in enumerate(hdr):
        set_cell_width(cell, widths[idx])
        set_cell_shading(cell, LIGHT_FILL)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                set_east_asia_font(run)
                run.font.bold = True
    for number, question in questions:
        cells = table.add_row().cells
        cells[0].text = str(number)
        cells[1].text = question
        set_cell_width(cells[0], widths[0])
        set_cell_width(cells[1], widths[1])
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for cell in cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_east_asia_font(run)
                    run.font.size = Pt(9.5)


def build_doc():
    doc = Document()
    configure_styles(doc)
    add_title(doc)

    raw_lines = SOURCE.read_text(encoding="utf-8").splitlines()
    questions = []
    for line in raw_lines:
        match = re.match(r"##\s+(\d+)\.\s+(.+)", line)
        if match:
            questions.append((int(match.group(1)), match.group(2).strip()))
    add_question_index(doc, questions)

    in_code = False
    code_lines = []
    skip_title = True
    for line in raw_lines:
        if skip_title and line.startswith("# "):
            skip_title = False
            continue

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue

        if in_code:
            code_lines.append(line)
            continue

        stripped = line.strip()
        if not stripped:
            continue

        h2 = re.match(r"##\s+(.+)", stripped)
        h3 = re.match(r"###\s+(.+)", stripped)
        bullet = re.match(r"-\s+(.+)", stripped)
        number = re.match(r"\d+\.\s+(.+)", stripped)

        if h2:
            doc.add_heading(h2.group(1), level=1)
        elif h3:
            doc.add_heading(h3.group(1), level=2)
        elif bullet:
            add_bullet(doc, bullet.group(1))
        elif number:
            add_numbered(doc, number.group(1))
        else:
            paragraph = doc.add_paragraph()
            parse_inline_code(paragraph, stripped)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
